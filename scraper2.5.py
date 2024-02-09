import os
import requests
from urllib.parse import urljoin, urlparse  # Added urlparse here
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO

class WebScraper:
    def __init__(self, base_url, output_file):
        self.base_url = base_url
        self.visited_urls = set()
        self.document = Document()
        self.output_file = output_file

    def get_page_content(self, url):
        try:
            response = requests.get(url)
            if response.status_code == 200:
                return response.text
            else:
                print(f"Error fetching {url}: Status code {response.status_code}")
                return None
        except Exception as e:
            print(f"An error occurred: {e}")
            return None

    def save_image(self, url):
        try:
            response = requests.get(url)
            if response.status_code == 200:
                image_stream = BytesIO(response.content)
                image = self.document.add_picture(image_stream)
                # Set width for the image (maintain aspect ratio)
                image.width = Inches(4)
                return image
            else:
                print(f"Error fetching image {url}: Status code {response.status_code}")
                return None
        except Exception as e:
            print(f"An error occurred while fetching image {url}: {e}")
            return None

    def parse_content(self, html, url):
        soup = BeautifulSoup(html, 'html.parser')

        # Add URL as a header
        self.document.add_heading('URL: ' + url, level=1)

        # Add the text content
        for paragraph in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            if paragraph.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self.document.add_heading(paragraph.get_text(), level=int(paragraph.name[1]))
            else:
                self.document.add_paragraph(paragraph.get_text())

        # Add hyperlinks (as plain text)
        for a_tag in soup.find_all('a', href=True):
            self.document.add_paragraph(a_tag['href'])

    def is_valid_url(self, url):
        parsed_url = urlparse(url)
        return bool(parsed_url.netloc) and parsed_url.netloc in self.base_url

    def scrape(self, url):
        if url not in self.visited_urls and self.is_valid_url(url):
            print(f"Scraping {url}")
            self.visited_urls.add(url)
            html_content = self.get_page_content(url)
            
            if html_content:
                self.parse_content(html_content, url)

                # Find all links in the page and recursively scrape them
                soup = BeautifulSoup(html_content, 'html.parser')
                for link in soup.find_all('a', href=True):
                    absolute_link = urljoin(url, link['href'])
                    self.scrape(absolute_link)

    def save_document(self):
        self.document.save(self.output_file)

def main():
    url = 'https://qo-docs.cambridgequantum.com/index.html'
    output_file = 'output2.docx'
    scraper = WebScraper(url, output_file)
    scraper.scrape(url)
    scraper.save_document()

if __name__ == "__main__":
    main()